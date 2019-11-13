# -*- coding: utf-8 -*-
"""
Created on Sat Aug 17 15:10:59 2019

@author: sairamtvv
"""

import numpy as np
import pandas as pd
import matplotlib.pyplot as plt
import pathlib
import docx
from docx.shared import Cm, Inches,Length

b = np.array( [[-372.675,372.434],[-372.671,372.437],[-372.669,372.435],
  [-372.665,372.435],[-372.666,372.436],[-372.665,372.439]])

#print(b.shape)

def matlab_conver_func (a,dir_for_analy,sensor_name):
    
    
    
    #SF inrun stability
    
    g=9.78335;
    R=200;
    
    
    U1=a[0][0] 
    U2=a[0][1];
    U3=a[1][0]; 
    U4=a[1][1];
    U5=a[2][0]; 
    U6=a[2][1];
    U7=a[3][0]; 
    U8=a[3][1];
    U9=a[4][0]; 
    U10=a[4][1];
    U11=a[5][0]; 
    U12=a[5][1];
    
    
    
    
    SF1=(-U1+U2)/(2*g*R); 
    SF2=(-U3+U4)/(2*g*R);
    SF3=(-U5+U6)/(2*g*R); 
    SF4=(-U7+U8)/(2*g*R);
    SF5=(-U9+U10)/(2*g*R); 
    SF6=(-U11+U12)/(2*g*R);
    
    #STD=std([SF1 SF2 SF3 SF4 SF5 SF6]);
    listSF=[SF1,SF2, SF3, SF4, SF5,SF6];
    STD=np.std(listSF);
    
    #MEAN=mean([SF1 SF2 SF3 SF4 SF5 SF6]);
    MEAN=np.mean(listSF);
    
    
    #Bias inrun stability
    a1=(a[1-1][1-1]+a[1-1][2-1])/(a[1-1][1-1]-a[1-1][2-1]);
    a2=(a[2-1][1-1]+a[2-1][2-1])/(a[2-1][1-1]-a[2-1][2-1]);
    a3=(a[3-1][1-1]+a[3-1][2-1])/(a[3-1][1-1]-a[3-1][2-1]);
    a4=(a[4-1][1-1]+a[4-1][2-1])/(a[4-1][1-1]-a[4-1][2-1]);
    a5=(a[5-1][1-1]+a[5-1][2-1])/(a[5-1][1-1]-a[5-1][2-1]);
    a6=(a[6-1][1-1]+a[6-1][2-1])/(a[6-1][1-1]-a[6-1][2-1]);
    
    SF_INRUN_STABILITY=(STD/MEAN)*10**6;
    #BIAS_INRUN_STABILITY=std([a1,a2,a3,a4,a5,a6])*10^6;
    BIAS_INRUN_STABILITY=np.std([a1,a2,a3,a4,a5,a6])*10**6;
    
    
    
    
    
    #SFTC
    
    #a1=load('RESULT1.dat');
    
    panda_a1 = pd.read_csv(dir_for_analy.joinpath('RESULT1.DAT'),header=None,delimiter= '\s+',skiprows=1)
    array_a1 = np.asarray(panda_a1)
    print(np.size(array_a1))
    a1=array_a1.reshape(9,5)
    
    
    #a2=load('RESULT2.dat');
    panda_a2 = pd.read_csv(dir_for_analy.joinpath('RESULT2.DAT'),header=None,delimiter= '\s+',skiprows=1)
    array_a2 = np.asarray(panda_a2)
    a2=array_a2.reshape(9,5)
    
    
    
    #a3=load('RESULT3.dat');
    panda_a3 = pd.read_csv(dir_for_analy.joinpath('RESULT3.DAT'),header=None,delimiter= '\s+',skiprows=1)
    array_a3 = np.asarray(panda_a3)
    a3=array_a3.reshape(9,5)
    
    
    #loop 1
    #%S.F(at +42.5°)
    #format short
    SF1=a1[2-1][2-1]*10;
    #SF TEMP.COEFFICENT
    #format long
    #70 deg C
    SFTC1=a1[1-1][2-1];
    SFTC2=a1[9-1][2-1];
    b1=np.mean([SFTC1,SFTC2]);
    #15 deg C
    SFTC3=a1[3-1][2-1];
    SFTC4=a1[7-1][2-1];
    b2=np.mean([SFTC3,SFTC4]);
    b3=a1[5-1][2-1];
    #format short
    loop1=((b1-b3)/b2)/110*10**6;
    
    
    
    
    #loop 2
    #S.F(at +42.5°)
    #format short
    SF2=a2[2-1][2-1]*10;
    #format long
    #SF TEMP.COEFFICENT
    #70 deg C
    SFTC1=a2[1-1][2-1];
    SFTC2=a2[9-1][2-1];
    b1=np.mean([SFTC1,SFTC2]);
    #15 deg C
    SFTC3=a2[3-1][2-1];
    SFTC4=a2[7-1][2-1];
    b2=np.mean([SFTC3,SFTC4]);
    #b3=a2(5,2)
    #format short
    loop2=((b1-b3)/b2)/110*10**6;
    
    
    
    #loop 3
    #S.F(at +42.5°)
    #format short
    SF3=a2[2-1][2-1]*10;
    #format long
    #SF TEMP.COEFFICENT
    #%70 deg C
    SFTC1=a3[1-1][2-1];
    SFTC2=a3[9-1][2-1];
    b1=np.mean([SFTC1,SFTC2]);
    #%15 deg C
    SFTC3=a3[3-1][2-1];
    SFTC4=a3[7-1][2-1];
    b2=np.mean([SFTC3,SFTC4]);
    #%b3=a2(5,2)
    #format short
    loop3=((b1-b3)/b2)/110*10**6;
    xsf = [SF1 ,SF2, SF3];
    SFAve=np.mean(xsf);
    Ysf=[loop1, loop2, loop3];
    SFTCAve=np.mean(Ysf); 
    
    
    
    
    #%BTCF
    
    
    #loop 1
    #%BIAS (at +42.5°)
    #format long
    B1=abs(a1[2-1][3-1]);
    B2=abs(a1[8-1][3-1]);
    #format short
    B3=(B1+B2)/2*10**6;
    #format long
    BTC1=a1[1-1][3-1];
    BTC2=a1[9-1][3-1];
    mBTC=np.mean([BTC1,BTC2]);
    BTC3=a1[5-1][3-1];
    #format short
    BTCF1=(mBTC-BTC3)/110*10**6;
    
    
    #loop2
    B1=abs(a2[2-1][3-1]);
    B2=abs(a2[8-1][3-1]);
    #format short
    B4=(B1+B2)/2*10**6;
    #format long
    BTC1=a2[1-1][3-1];
    BTC2=a2[9-1][3-1];
    mBTC=np.mean([BTC1, BTC2]);
    BTC3=a2[5-1][3-1];
    #format short
    BTCF2=(mBTC-BTC3)/110*10**6;
    
    
    #loop3
    B1=abs(a3[2-1][3-1]);
    B2=abs(a3[8-1][3-1]);
    #format short
    B5=(B1+B2)/2*10**6;
    #format long
    BTC1=a3[1-1][3-1];
    BTC2=a3[9-1][3-1];
    mBTC=np.mean([BTC1,BTC2]);
    BTC3=a3[5-1][3-1];
    #format short
    BTCF3=(mBTC-BTC3)/110*10**6;
    xb = [ B3, B4, B5];
    BiasAve=np.mean(xb);
    yb=[BTCF1, BTCF2, BTCF3];
    BTCAve=np.mean(yb);
    
    
    #%SF HYSTERESIS & BIAS HYSTERESIS ERROR
    
    b1=a1[1-1][2-1];
    b2=a1[9-1][2-1];
    b12=np.mean([b1, b2]);
    Hyst19=np.diff([b1, b2])/b12;
    b3=a1[2-1][2-1];
    b4=a1[8-1][2-1];
    b34=np.mean([b3, b4]);
    Hyst28=np.diff([b3, b4])/b34;
    b5=a1[3-1][2-1];
    b6=a1[7-1][2-1];
    b56=np.mean([b5, b6]);
    Hyst37=np.diff([b5, b6])/b56;
    b7=a1[4-1][2-1];
    b8=a1[6-1][2-1];
    b78=np.mean([b7, b8]);
    Hyst46=np.diff([b7, b8])/b78;
    c=np.abs([Hyst19, Hyst28, Hyst37, Hyst46]);
    SF_HYSTERESIS1=np.max(c)*10**6;
    #loop2
    b1=a2[1-1][2-1];
    b2=a2[9-1][2-1];
    b12=np.mean([b1, b2]);
    Hyst19=np.diff([b1, b2])/b12;
    b3=a2[2-1][2-1];
    b4=a2[8-1][2-1];
    b34=np.mean([b3, b4]);
    Hyst28=np.diff([b3, b4])/b34;
    b5=a2[3-1][2-1];
    b6=a2[7-1][2-1];
    b56=np.mean([b5, b6]);
    Hyst37=np.diff([b5, b6])/b56;
    b7=a2[4-1][2-1];
    b8=a2[6-1][2-1];
    b78=np.mean([b7, b8]);
    Hyst46=np.diff([b7, b8])/b78;
    c=np.abs([Hyst19, Hyst28, Hyst37, Hyst46]);
    SF_HYSTERESIS2=np.max(c)*10**6;
    #loop3
    b1=a3[1-1][2-1];
    b2=a3[9-1][2-1];
    b12=np.mean([b1, b2]);
    Hyst19=np.diff([b1, b2])/b12;
    b3=a3[2-1][2-1];
    b4=a3[8-1][2-1];
    b34=np.mean([b3, b4]);
    Hyst28=np.diff([b3, b4])/b34;
    b5=a3[3-1][2-1];
    b6=a3[7-1][2-1];
    b56=np.mean([b5, b6]);
    Hyst37=np.diff([b5, b6])/b56;
    b7=a3[4-1][2-1];
    b8=a3[6-1][2-1];
    b78=np.mean([b7, b8]);
    Hyst46=np.diff([b7, b8])/b78;
    c=np.abs([Hyst19, Hyst28, Hyst37, Hyst46]);
    SF_HYSTERESIS3=np.max(c)*10**6;
    #format short
    SF_HYSTERESIS=([SF_HYSTERESIS1,  SF_HYSTERESIS2,   SF_HYSTERESIS3]);
    avg_SF_HYSTERESIS_value=np.mean([SF_HYSTERESIS1, SF_HYSTERESIS2,   SF_HYSTERESIS3]);
    
    
    #BIAS HYSTERESIS ERROR
    b1=a1[1-1][3-1];
    b2=a1[9-1][3-1];
    BIAS_Hyst19=np.diff([b1, b2]);
    b3=a1[2-1][3-1];
    b4=a1[8-1][3-1];
    BIAS_Hyst28=np.diff([b3, b4]);
    b5=a1[3-1][3-1];
    b6=a1[7-1][3-1];
    BIAS_Hyst37=np.diff([b5, b6]);
    b7=a1[4-1][3-1];
    b8=a1[6-1][3-1];
    BIAS_Hyst46=np.diff([b7, b8]);
    c=np.abs([BIAS_Hyst19, BIAS_Hyst28, BIAS_Hyst37, BIAS_Hyst46]);
    BIAS_HYSTERESIS1=np.max(c)*10**6;
    #%loop2
    b1=a2[1-1][3-1];
    b2=a2[9-1][3-1];
    BIAS_Hyst19=np.diff([b1, b2]);
    b3=a2[2-1][3-1];
    b4=a2[8-1][3-1];
    BIAS_Hyst28=np.diff([b3, b4]);
    b5=a2[3-1][3-1];
    b6=a2[7-1][3-1];
    BIAS_Hyst37=np.diff([b5, b6]);
    b7=a2[4-1][3-1];
    b8=a2[6-1][3-1];
    BIAS_Hyst46=np.diff([b7, b8]);
    c=np.abs([BIAS_Hyst19, BIAS_Hyst28, BIAS_Hyst37, BIAS_Hyst46]);
    BIAS_HYSTERESIS2=np.max(c)*10**6;
    #% %%loop3
    b1=a3[1-1][3-1];
    b2=a3[9-1][3-1];
    BIAS_Hyst19=np.diff([b1, b2]);
    b3=a3[2-1][3-1];
    b4=a3[8-1][3-1];
    BIAS_Hyst28=np.diff([b3, b4]);
    b5=a3[3-1][3-1];
    b6=a3[7-1][3-1];
    BIAS_Hyst37=np.diff([b5, b6]);
    b7=a3[4-1][3-1];
    b8=a3[6-1][3-1];
    BIAS_Hyst46=np.diff([b7, b8]);
    c=np.abs([BIAS_Hyst19, BIAS_Hyst28, BIAS_Hyst37, BIAS_Hyst46]);
    BIAS_HYSTERESIS3=np.max(c)*10**6;
    #format short
    BIAS_HYSTERESIS=([BIAS_HYSTERESIS1, BIAS_HYSTERESIS2, BIAS_HYSTERESIS3]);
    avg_BIAS_HYSTERESIS_value=np.mean([BIAS_HYSTERESIS1, BIAS_HYSTERESIS2, BIAS_HYSTERESIS3]);
    
    
    
    
    
    
    
    
    
    
    
    #%S.F Day to Day Stability (at +42.5°)
    #format long
    SF1=a1[2-1][2-1]*10;
    SF2=a1[8-1][2-1]*10;
    SF12=np.mean([SF1, SF2]);
    SF3=a2[2-1][2-1]*10;
    SF4=a2[8-1][2-1]*10;
    SF34=np.mean([SF3, SF4]);
    SF5=a3[2-1][2-1]*10;
    SF6=a3[8-1][2-1]*10;
    SF56=np.mean([SF5, SF6]);
    SF7=np.mean([SF1, SF2, SF3, SF4,SF5,SF6]);
    SFDD1=abs((SF12-SF34)/SF7)*10**6;
    SFDD2=abs((SF34-SF56)/SF7)*10**6;
    SFDD3=np.mean([SFDD1, SFDD2]);
    xsfdd=[SFDD1, SFDD2, SFDD3];
    
    #%BIAS Day to Day Stability (at +42.5°)
    #format long
    B1=a1[2-1][3-1];
    B2=a1[8-1][3-1];
    B12=np.mean([B1, B2])*10**6;
    B3=a2[2-1][3-1];
    B4=a2[8-1][3-1];
    B34=np.mean([B3, B4])*10**6;
    B5=a3[2-1][3-1];
    B6=a3[8-1][3-1];
    B56=np.mean([B5, B6])*10**6;
    BDD1=abs(B12-B34);
    BDD2=abs(B34-B56);
    BDD3=np.mean([BDD1, BDD2]);
    ybdd=[BDD1, BDD2, BDD3];
    #format short
    
    
    
    
    #%MISALIGNMENT
    
    #%loop1
    #[C1,I1]=max(abs(a1(:,4)));
    I1 =  np.unravel_index(np.argmax(np.abs(a1[:,4-1]), axis=None), np.abs(a1[:,4-1]).shape)[0]
    MIS1=a1[I1][4-1];
     
    #%loop2
    I2 =  np.unravel_index(np.argmax(np.abs(a2[:,4-1]), axis=None), np.abs(a2[:,4-1]).shape)[0]
    #[C2,I2]=max(abs(a2(:,4)));
    MIS2=a2[I2][4-1];
    
    #%loop3
    #[c3,I3]=max(abs(a3(:,4)));
    I3 =  np.unravel_index(np.argmax(np.abs(a3[:,4-1]), axis=None), np.abs(a3[:,4-1]).shape)[0]
    MIS3=a3[I3][4-1];
    
    xm = [ MIS1, MIS2, MIS3];
    
    #xmm=xm';
    xmm=np.transpose(xm);
    
    #[C4,I4]=max(abs(xm));
    I4 =  np.unravel_index(np.argmax(np.abs(xm), axis=None), np.abs(xm).shape)[0]
    
    #M_max=xmm(I4,1);
    M_max=xmm[I4];
    
    
    
    #%Misalignment day to day
    
    m1=a1[:,4-1];
    m2=a2[:,4-1];
    m3=a3[:,4-1];
    md1=np.abs(m1-m2);
    md2=np.abs(m2-m3);
    md2d=[np.max(md1), np.max(md2)];
    avge=np.mean(md2d);
    
    #format short
    
    
    print('     ');
    #print("Sammy ate {0:} percent of a pizza!".format(xsf))
    print('              Loop1           Loop2           Loop3           Average');
    print('     ');
    print('SF:         {0:6.2f}          {1:6.2f}           {2:6.2f}          {3:6.2f}'.format(xsf[0],xsf[1],xsf[2],np.mean(xsf)))
    print('     ');
    print('SFTC:       {0:6.2f}          {1:6.2f}           {2:6.2f}          {3:6.2f}'.format(Ysf[0],Ysf[1],Ysf[2],np.mean(Ysf)))
    print('     ');
    print('SF_HYS:     {0:6.2f}          {1:6.2f}           {2:6.2f}          {3:6.2f}'.format(SF_HYSTERESIS[0],SF_HYSTERESIS[1],SF_HYSTERESIS[2],avg_SF_HYSTERESIS_value))
    print('     ');
    print('SFDD:       {0:6.2f}          {1:6.2f}           {2:6.2f}                  '.format(xsfdd[0],xsfdd[1],xsfdd[2]))
    print('     ');
    print('SF_INRUN_S::--------          -------           --------    {0:6.2f}'.format(SF_INRUN_STABILITY))
    print('     ');
    print('Bias:       {0:6.2f}          {1:6.2f}           {2:6.2f}          {3:6.2f}'.format(xb[0],xb[1],xb[2],np.mean(xb)))
    print('     ');
    print('BTCF:       {0:6.2f}          {1:6.2f}           {2:6.2f}          {3:6.2f}'.format(yb[0],yb[1],yb[2],np.mean(yb)))
    print('     ');
    print('BIAS_HYS:   {0:6.2f}          {1:6.2f}           {2:6.2f}          {3:6.2f}'.format(BIAS_HYSTERESIS[0],BIAS_HYSTERESIS[1],BIAS_HYSTERESIS[2],avg_BIAS_HYSTERESIS_value))
    print('     ');
    print('BiasDD:     --------          {0:6.2f}           {1:6.2f}      {2:6.2f}'.format(ybdd[0],ybdd[1],ybdd[2]))
    print('     ');
    print('B_INRUN_S:  --------          --------           --------    {0:6.2f}'.format(BIAS_INRUN_STABILITY))
    print('     ');
    print('MIS:        {0:6.2f}          {1:6.2f}           {2:6.2f}          {3:6.2f}'.format(xm[0],xm[1],xm[2],M_max))
    print('     ');
    print('MD2D:       ---------         {0:6.2f}           {1:6.2f}          {2:6.2f}'.format(md2d[0],md2d[1],np.max(md2d)))
    
    
    
    #df = pd.DataFrame(data)
    df = pd.read_csv(dir_for_analy.joinpath("sample1.csv"))
    
    #Assigning the values into the dataframe 
    df.iloc[10,3]=np.mean(xsf) #sf mean
    df.iloc[11,3]=np.mean(Ysf) #sftc mean
    df.iloc[12,3]= avg_SF_HYSTERESIS_value #sf_hysterisis
    df.iloc[13,3]= xsfdd[2] #SFDD
    df.iloc[14,3]= SF_INRUN_STABILITY #SF_INRUN_S
    df.iloc[15,3]= np.mean(xb)#bias
    df.iloc[16,3]= np.mean(yb)#BTCF bias temperature coefficient
    df.iloc[17,3]= avg_BIAS_HYSTERESIS_value #BIAS_HYS
    df.iloc[18,3]= ybdd[2] #BiasDD Bias Day to Day Stability 	
    df.iloc[19,3]= BIAS_INRUN_STABILITY        #B_INRUN_S
    df.iloc[20,3]= M_max #MIS  Misalignment 	
    df.iloc[21,3]= np.max(md2d)    #MD2D Axis Misalignment Stability (Day to Day)
    
    
    
    
    
    # open an existing document
    doc = docx.Document()
    section = doc.sections[0]
    header = section.header
    footer = section.footer
    header_p = header.paragraphs[0]
    footer_p=footer.paragraphs[0]
    header_p.text = "EQA3/AT/QC/QCR/03"
    header_p.alignment =  2
    header_p.bold= True
    
    footer_p.text="EQA3/AT/QC/QCR/03Revision: 00      Date:"
    footer_p.alignment =  1
    footer_p.bold= True
    
    
    doc.add_heading('QUALITY CONFORMANCE REPORT - ACCEPTANCE TESTS ', 1)
    
    
    #project_p = doc.add_paragraph("Date:")
    #project_p.alignment = 1 # for left, 1 for right, 2 center, 3 justify ....
    #project_p.bold = True
    
    
    
    date_p= doc.add_paragraph("PROJECT: EQA-3                \t\t\t        Date: ")
    date_p.alignment = 1 # for left, 1 for right, 2 center, 3 justify ....
    date_p.bold = True
    
    
    # add a table to the end and create a reference variable
    # extra row is so we can add the header row
    table = doc.add_table(df.shape[0]+1, df.shape[1])
    
    
    table.style = 'TableGrid' #single lines in all cells
    table.autofit = False
    
    col = table.columns[0] 
    col.width=Inches(0.75)
    cell=table.cell(1,1)
    cell.width = Inches(2)
    #col.width=Cm(1.0)
    #col.width=360000 #=1cm
    #for cell in table.cells:
    #    cell.width = Inches(1)
    
    
    
    
    print(table)
    print('------------------------------------------')
    
    # add the header rows.
    for j in range(df.shape[-1]):
        if df.columns[j] == '*':
            pass
        else:
            table.cell(0,j).text = df.columns[j]
    
    # add the rest of the data frame
    for i in range(df.shape[0]):
        for j in range(df.shape[-1]):
            if df.values[i,j] == '*':
                pass
            else:
                table.cell(i+1,j).text = str(df.values[i,j])
    
            #table.cell(i+1,j).text = str(df.values[i,j])
            
    note_p= doc.add_paragraph("Note:**The test results are calculated from the data of ATLs’-14, 16, 17 and 20. ")
    note_p.alignment = 1 # for left, 1 for right, 2 center, 3 justify ....
    note_p.bold = True        
    
    
    doc.add_page_break()
    
    
    
    df = pd.read_csv(dir_for_analy.joinpath("sample2.csv"))
    table = doc.add_table(df.shape[0]+1, df.shape[1])
    
    
    table.style = 'TableGrid' #single lines in all cells
    table.autofit = False
    
    col = table.columns[0] 
    col.width=Inches(0.75)
    cell=table.cell(1,1)
    cell.width = Inches(2)
    
    print(table)
    print('------------------------------------------')
    
    # add the header rows.
    for j in range(df.shape[-1]):
        if df.columns[j] == '*':
            pass
        else:
            table.cell(0,j).text = df.columns[j]
    
    # add the rest of the data frame
    for i in range(df.shape[0]):
        for j in range(df.shape[-1]):
            if df.values[i,j] == '*':
                pass
            else:
                table.cell(i+1,j).text = str(df.values[i,j])
    
    # save the doc
    doc.save(dir_for_analy.joinpath('conform_'+sensor_name+'.docx'))
    
    plotting_graph(dir_for_analy,sensor_name)
    
    
def endurance_test(df_endu_channels,dir_for_analy,sensor_name):
    
    doc = docx.Document()
    section = doc.sections[0]
    header = section.header
    footer = section.footer
    header_p = header.paragraphs[0]
    footer_p=footer.paragraphs[0]
    header_p.text = "EQA3/ELEC/ET/QC/ATL/19"
    header_p.alignment =  2
    header_p.bold= True
    
    footer_p.text="EQA3/ELEC/ET/QC/ATL/19 \t Revision: 00 \t Date:"
    footer_p.alignment =  0
    footer_p.bold= True
    
    
    doc.add_heading('ACCEPTANCE TEST LEAF FOR ENDURANCE TEST ', 1)
    
    
    #project_p = doc.add_paragraph("Date:")
    #project_p.alignment = 1 # for left, 1 for right, 2 center, 3 justify ....
    #project_p.bold = True
    string_project="PROJECT: EQA-3 \nEQA-3 NO.__________________________\n"\
                    +"Start Date:_________ \t \t \t \t \t \t "\
                    +"END Date:___________ \n"\
                    +'Test Condition:"+1g" at Ambient \n'\
                    +"Test Reading Interval every 2 Hrs\n"\
                    +"Total Endurance Test Time 24 Hrs"
                    
    
    
    project_p= doc.add_paragraph(string_project)
    project_p.alignment = 0 # for left, 1 for right, 2 center, 3 justify ....
    project_p.bold = True
    
    
    df = pd.read_csv(dir_for_analy.joinpath("endu_sample1.csv"))
    
    
    # add a table to the end and create a reference variable
    # extra row is so we can add the header row
    table = doc.add_table(df.shape[0]+1, df.shape[1])
    
    
    table.style = 'TableGrid' #single lines in all cells
    table.autofit = False
    
    col = table.columns[0] 
    col.width=Inches(0.75)
    cell=table.cell(1,1)
    cell.width = Inches(2)
    #col.width=Cm(1.0)
    #col.width=360000 #=1cm
    #for cell in table.cells:
    #    cell.width = Inches(1)
    
    
    
    
    print(table)
    print('------------------------------------------')
    
    # add the header rows.
    for j in range(df.shape[-1]):
        if df.columns[j] == '*':
            pass
        else:
            table.cell(0,j).text = df.columns[j]
    
    # add the rest of the data frame
    for i in range(df.shape[0]):
        for j in range(df.shape[-1]):
            if df.values[i,j] == '*':
                pass
            else:
                table.cell(i+1,j).text = str(df.values[i,j])
    
            #table.cell(i+1,j).text = str(df.values[i,j])
    
    #Finished table 1 which is actually one line
    #-----------------------------------------------------------------
    newline_p= doc.add_paragraph("\n\n")
    
    df = pd.read_csv(dir_for_analy.joinpath("endu_sample2.csv"))
    
    for i,j in zip(list(range(0,12)),list(range(0,2))):
        df.iloc[i,j+1]=df_endu_channels.iloc[i,j]
    
    
    
    
    # add a table to the end and create a reference variable
    # extra row is so we can add the header row
    table = doc.add_table(df.shape[0]+1, df.shape[1])
    
    
    table.style = 'TableGrid' #single lines in all cells
    table.autofit = False
    
    col = table.columns[0] 
    col.width=Inches(0.75)
    cell=table.cell(1,1)
    cell.width = Inches(2)
    #col.width=Cm(1.0)
    #col.width=360000 #=1cm
    #for cell in table.cells:
    #    cell.width = Inches(1)
    
    
    
    
    print(table)
    print('------------------------------------------')
    
    # add the header rows.
    for j in range(df.shape[-1]):
        if df.columns[j] == '*':
            pass
        else:
            table.cell(0,j).text = df.columns[j]
    
    # add the rest of the data frame
    for i in range(df.shape[0]):
        for j in range(df.shape[-1]):
            if df.values[i,j] == '*':
                pass
            else:
                table.cell(i+1,j).text = str(df.values[i,j])
                #table.cell(i+1,j).text = str(df.values[i,j])
    #end of table 2            
    #-------------------------------------------------------------------------
    newline2_p= doc.add_paragraph("\n\n")        
    string_note="Note:**The above readings donot have any functional significance. The test is\n"\
                +'only to check the Survivability of Sensor in Power "ON" condition' 
    note_p= doc.add_paragraph(string_note)
    note_p.alignment = 0 # for left, 1 for right, 2 center, 3 justify ....
    note_p.bold = True  
    
    #---------------------------------------------------
    #start of table 3
    df = pd.read_csv(dir_for_analy.joinpath("endu_sample3.csv"))
    
    
    # add a table to the end and create a reference variable
    # extra row is so we can add the header row
    table = doc.add_table(df.shape[0]+1, df.shape[1])
    
    
    table.style = 'TableGrid' #single lines in all cells
    table.autofit = False
    
    col = table.columns[0] 
    col.width=Inches(0.75)
    cell=table.cell(1,1)
    cell.width = Inches(2)
    #col.width=Cm(1.0)
    #col.width=360000 #=1cm
    #for cell in table.cells:
    #    cell.width = Inches(1)
    
    
    
    
    print(table)
    print('------------------------------------------')
    
    # add the header rows.
    for j in range(df.shape[-1]):
        if df.columns[j] == '*':
            pass
        else:
            table.cell(0,j).text = df.columns[j]
    
    # add the rest of the data frame
    for i in range(df.shape[0]):
        for j in range(df.shape[-1]):
            if df.values[i,j] == '*':
                pass
            else:
                table.cell(i+1,j).text = str(df.values[i,j])
                #table.cell(i+1,j).text = str(df.values[i,j])
    #end of table 2      
    doc.save(dir_for_analy.joinpath('endurance_'+sensor_name+'.docx'))
        
    



    
def plotting_graph(dir_for_analy,sensor_name):
    
    #Getting the dates of their execution
    df_day1=pd.read_csv(dir_for_analy.parent.joinpath('raw_text_L1.txt'))
    label1=df_day1.iloc[0,1].split()[0]
    df_day2=pd.read_csv(dir_for_analy.parent.joinpath('raw_text_L2.txt'))
    label2=df_day2.iloc[0,1].split()[0]
    df_day3=pd.read_csv(dir_for_analy.parent.joinpath('raw_text_L3.txt'))
    label3=df_day3.iloc[0,1].split()[0]
    
    #Plotting the graphs
    pdresult1 = pd.read_csv(dir_for_analy.joinpath('RESULT1.DAT'),delimiter= '\s+')
    pdresult2 = pd.read_csv(dir_for_analy.joinpath('RESULT2.DAT'),delimiter= '\s+')
    pdresult3 = pd.read_csv(dir_for_analy.joinpath('RESULT3.DAT'),delimiter= '\s+')
    
    #plt.rc('text', usetex=True)
    
    x = list(pdresult1['Ts'])
    y_list1 = list(pdresult1['Misalign34'])
    y_list2 = list(pdresult2['Misalign34'])
    y_list3 = list(pdresult3['Misalign34'])
    
    # Create Figure (empty canvas)
    fig = plt.figure()
    
    # Add set of axes to figure
    axes = fig.add_axes([0.1, 0.1, 1, 1]) # left, bottom, width, height (range 0 to 1)
    
    # Plot on that set of axes
    axes.grid(True)
    axes.plot(x, y_list1,color='red',linewidth=1,alpha=1,linestyle='-',marker='o',markersize=2,
              markerfacecolor='red', markeredgewidth=3, markeredgecolor='red',label=label1)
    
    
    axes.plot(x, y_list2,color='green',linewidth=1,alpha=1,linestyle='-',marker='s',markersize=2,
              markerfacecolor='green', markeredgewidth=3, markeredgecolor='green',label=label2)
    axes.plot(x, y_list3,color='blue',linewidth=1,alpha=1,linestyle='-',marker='s',markersize=2,
              markerfacecolor='blue', markeredgewidth=3, markeredgecolor='blue',label=label3)
    
    #axes.plot(x,z)
    axes.set_xlabel('Temp $^\circ$ C') # Notice the use of set_ to begin methods  
    axes.set_ylabel('Scale, $mA(m/s^2)$')
    #axes.set_title('Set Title')
    axes.legend(loc=0)
    # legend
    #plt.legend(('phase field', 'level set', 'sharp interface'),
    #           shadow=True, loc=(0.01, 0.48), handlelength=1.5, fontsize=16)
    
    
    #fig.savefig("filename.png", dpi=200)
    
    locs, labels = plt.yticks() 
    for index,tuplexy in enumerate(zip(x, y_list1)):
        i_x=tuplexy[0]
        i_y=tuplexy[1]
        #print("{},{},{}".format(index,i_x, i_y))
        label='{1:.7f}[{0}]'.format(index+1, i_y)
        #plt.text(i_x, i_y,string )
        plt.annotate(label, # this is the text
                     (i_x,i_y), # this is the point to label
                     textcoords="offset points", # how to position the text
                     xytext=(0,-2*index+20), # distance from text to points (x,y)
                     ha='center',color='red') # horizontal alignment can be left, right or center
        
        
    for index,tuplexy in enumerate(zip(x, y_list2)):
        i_x=tuplexy[0]
        i_y=tuplexy[1]
        #print("{},{},{}".format(index,i_x, i_y))
        label='{1:.7f}[{0}]'.format(index+1, i_y)
        #plt.text(i_x, i_y,string )
        plt.annotate(label, # this is the text
                     (i_x,i_y), # this is the point to label
                     textcoords="offset points", # how to position the text
                     xytext=(0,-3*index-20), # distance from text to points (x,y)
                     ha='center',color='green') # horizontal alignment can be left, right or center   
    
    for index,tuplexy in enumerate(zip(x, y_list3)):
        i_x=tuplexy[0]
        i_y=tuplexy[1]
        #print("{},{},{}".format(index,i_x, i_y))
        label='{1:.7f}[{0}]'.format(index+1, i_y)
        #plt.text(i_x, i_y,string )
        plt.annotate(label, # this is the text
                     (i_x,i_y), # this is the point to label
                     textcoords="offset points", # how to position the text
                     xytext=(0,-3*index-60), # distance from text to points (x,y)
                     ha='center',color='blue') # horizontal alignment can be left, right or center 
    plt.savefig(dir_for_analy.joinpath('fig_'+sensor_name+'.png'))
    plt.show()


#
#
#
#
#
#
#
#
#
#x = [1,2,3,4]
#y = [1,4,9,16]
#z=[1,8,27,64]
## Create Figure (empty canvas)
#fig = plt.figure()
#
## Add set of axes to figure
#axes = fig.add_axes([0.1, 0.1, 0.8, 0.8]) # left, bottom, width, height (range 0 to 1)
#
## Plot on that set of axes
#axes.plot(x, y,color='purple',linewidth=2,alpha=1,linestyle='-',marker='o',markersize=7,
#          markerfacecolor='purple', markeredgewidth=3, markeredgecolor='purple')
#axes.plot(x,z)
#axes.set_xlabel('Set X Label') # Notice the use of set_ to begin methods  
#axes.set_ylabel('Set y Label')
#axes.set_title('Set Title')
#
#fig.savefig("filename.png", dpi=200)
#
#for i_x, i_y in zip(x, y):
#    plt.text(i_x, i_y, '({}, {})'.format(i_x, i_y))
#
#plt.show()
#
#
#
#
#
#
#
#
#
#
#
#
#
#
## writingto excel file for D4 and D5 analysis
#import openpyxl
#
#srcfile = openpyxl.load_workbook('EQA3-31_(AS)_adj_param_p20.xlsx',read_only=False, keep_vba= True)#to open the excel sheet and if it has macros
#
#sheetname = srcfile.get_sheet_by_name('sheetsai')#get sheetname from the file
#sheetname['C4']= 55.568 #write something in B2 cell of the supplied sheet
#sheetname.cell(row=1,column=1).value = "something" #write to row 1,col 1 explicitly, this type of writing is useful to write something in loops
#
#srcfile.save('newfile.xlsm')#save it as a new file, the original file is untouched and here I am saving it as xlsm(m here denotes macros).
#
#
#
#import pandas as pd
#pdresult1 = pd.read_csv('RESULT1.DAT',delimiter= '\s+')
#pdresult2 = pd.read_csv('RESULT2.DAT',delimiter= '\s+')
#pdresult3 = pd.read_csv('RESULT3.DAT',delimiter= '\s+')
#import matplotlib.pyplot as plt
##plt.rc('text', usetex=True)
#
#x = list(pdresult1['Ts'])
#y_list1 = list(pdresult1['Misalign34'])
#y_list2 = list(pdresult2['Misalign34'])
#y_list3 = list(pdresult3['Misalign34'])
#
## Create Figure (empty canvas)
#fig = plt.figure()
#
## Add set of axes to figure
#axes = fig.add_axes([0.1, 0.1, 1, 1]) # left, bottom, width, height (range 0 to 1)
#
## Plot on that set of axes
#axes.grid(True)
#axes.plot(x, y_list1,color='red',linewidth=1,alpha=1,linestyle='-',marker='o',markersize=2,
#          markerfacecolor='red', markeredgewidth=3, markeredgecolor='red',label="x**2")
#
#
#axes.plot(x, y_list2,color='green',linewidth=1,alpha=1,linestyle='-',marker='s',markersize=2,
#          markerfacecolor='green', markeredgewidth=3, markeredgecolor='green',label="x**3")
#axes.plot(x, y_list3,color='blue',linewidth=1,alpha=1,linestyle='-',marker='s',markersize=2,
#          markerfacecolor='blue', markeredgewidth=3, markeredgecolor='blue',label="x**4")
#
##axes.plot(x,z)
#axes.set_xlabel('Temp $^\circ$ C') # Notice the use of set_ to begin methods  
#axes.set_ylabel('Scale, $mA(m/s^2)$')
##axes.set_title('Set Title')
#axes.legend(loc=0)
## legend
##plt.legend(('phase field', 'level set', 'sharp interface'),
##           shadow=True, loc=(0.01, 0.48), handlelength=1.5, fontsize=16)
#
#
##fig.savefig("filename.png", dpi=200)
#
#locs, labels = plt.yticks() 
#for index,tuplexy in enumerate(zip(x, y_list1)):
#    i_x=tuplexy[0]
#    i_y=tuplexy[1]
#    #print("{},{},{}".format(index,i_x, i_y))
#    label='{1:.7f}[{0}]'.format(index+1, i_y)
#    #plt.text(i_x, i_y,string )
#    plt.annotate(label, # this is the text
#                 (i_x,i_y), # this is the point to label
#                 textcoords="offset points", # how to position the text
#                 xytext=(0,-2*index+20), # distance from text to points (x,y)
#                 ha='center',color='red') # horizontal alignment can be left, right or center
#    
#    
#for index,tuplexy in enumerate(zip(x, y_list2)):
#    i_x=tuplexy[0]
#    i_y=tuplexy[1]
#    #print("{},{},{}".format(index,i_x, i_y))
#    label='{1:.7f}[{0}]'.format(index+1, i_y)
#    #plt.text(i_x, i_y,string )
#    plt.annotate(label, # this is the text
#                 (i_x,i_y), # this is the point to label
#                 textcoords="offset points", # how to position the text
#                 xytext=(0,-3*index-20), # distance from text to points (x,y)
#                 ha='center',color='green') # horizontal alignment can be left, right or center   
#
#for index,tuplexy in enumerate(zip(x, y_list3)):
#    i_x=tuplexy[0]
#    i_y=tuplexy[1]
#    #print("{},{},{}".format(index,i_x, i_y))
#    label='{1:.7f}[{0}]'.format(index+1, i_y)
#    #plt.text(i_x, i_y,string )
#    plt.annotate(label, # this is the text
#                 (i_x,i_y), # this is the point to label
#                 textcoords="offset points", # how to position the text
#                 xytext=(0,-3*index-60), # distance from text to points (x,y)
#                 ha='center',color='blue') # horizontal alignment can be left, right or center       
#    
#
#autoit.win_wait_active("[TITLE:Set Data Log Fields]", "")
#autoit.control_click("[TITLE:Set Data Log Fields]","", "[CLASS::WindowsForms10.BUTTON.app.0.33c0d9d; INSTANCE:3]") 
#ControlClick(WinGetTitle("[active]"), "", "[CLASS:WindowsForms10.BUTTON.app.0.143a722_r15_ad1; INSTANCE:1]")
#
#
#
##click the two dots so that it opens the dialogue box
##wait for 16 seconds to make the Benchlink Data logeer 3 active
#autoit.win_wait_active("[TITLE:Configuration-2-BenchLink Data Logger3]", "",16)
##XXXXXXXXXXXXXXXXXXX
#autoit.control_click("[TITLE:Configuration-2-BenchLink Data Logger3]","", "[CLASS::WindowsForms10.BUTTON.app.0.33c0d9d; INSTANCE:15]") 
#time.sleep(0.5)
#
##Click on the check box 
#autoit.win_wait_active("[TITLE:Set Data Log Fields]", "",16)
#autoit.control_click("[TITLE:Set Data Log Fields]","", "[CLASS::WindowsForms10.BUTTON.app.0.33c0d9d; INSTANCE:3]") 
#time.sleep(0.5)
#
#
##Click on the ok box
#
#
#
#
##Click on the play button
#autoit.win_wait_active("[TITLE:Configuration-2-BenchLink Data Logger3]", "",16)
#
##wait for the scan and log data summary
##wait for 2 minutes maximum
#autoit.win_wait_active("[TITLE:Scan and Log Data Summary]", "",150)
#autoit.control_click("[TITLE:Scan and Log Data Summary]","", "[CLASS::WindowsForms10.BUTTON.app.0.33c0d9d; INSTANCE:5]") 
#time.sleep(0.5)
#
#
##Click on the close button 
#
#
#
##end of the cycle
#
#
#
#
#
#
#
#
#
#
#
